import io
from typing import List
from pypdf import PdfReader  # <-- Updated from PyPDF2 to pypdf
import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

def _process_single_file(file) -> List[Document]:
    """Extract text from a single uploaded file and return a list of Documents."""
    docs: List[Document] = []
    file_name = file.name
    file_ext = file_name.split(".")[-1].lower()

    if file_ext == "pdf":
        pdf_bytes = io.BytesIO(file.read())
        pdf_reader = PdfReader(pdf_bytes)
        for page_num, page in enumerate(pdf_reader.pages, start=1):
            text = page.extract_text()
            if text and text.strip():
                docs.append(
                    Document(page_content=text, metadata={"source": file_name, "page": page_num})
                )
    elif file_ext == "docx":
        doc_bytes = io.BytesIO(file.read())
        doc = docx.Document(doc_bytes)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if full_text:
            docs.append(Document(page_content=full_text, metadata={"source": file_name, "page": 1}))
    elif file_ext == "txt":
        text = file.read().decode("utf-8")
        if text.strip():
            docs.append(Document(page_content=text, metadata={"source": file_name, "page": 1}))
    return docs


def process_uploaded_files(uploaded_files, progress=None) -> List[Document]:
    """Process a list of uploaded files in parallel.
    Args:
        uploaded_files: Iterable of Streamlit uploaded file objects.
        progress: Optional Streamlit progress widget (st.progress) to update.
    Returns:
        Flattened, ordered list of Document objects.
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    # Submit all files to the thread pool
    with ThreadPoolExecutor() as executor:
        future_to_file = {executor.submit(_process_single_file, f): f for f in uploaded_files}
        all_docs: List[Document] = []
        completed = 0
        total = len(uploaded_files)
        for future in as_completed(future_to_file):
            docs = future.result()
            all_docs.extend(docs)
            completed += 1
            if progress is not None:
                progress.progress(completed / total)

    # Ensure deterministic ordering by source filename then page number
    all_docs.sort(key=lambda d: (d.metadata.get("source", ""), d.metadata.get("page", 0)))
    return all_docs

def get_text_chunks(documents: List[Document]) -> List[Document]:
    """Split documents into contextual chunks using RecursiveCharacterTextSplitter."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", " ", ""]
    )
    return text_splitter.split_documents(documents)